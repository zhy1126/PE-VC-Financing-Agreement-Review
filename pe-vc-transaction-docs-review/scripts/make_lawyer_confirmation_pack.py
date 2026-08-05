#!/usr/bin/env python3
"""Generate an immutable-JSON-backed Chinese lawyer confirmation Word pack."""

from __future__ import annotations

import argparse
from copy import deepcopy
import json
from pathlib import Path
import shutil
import sys
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from confirmation_word_common import (
    canonical_json_sha256,
    document_part_sha256,
    inspect_docx,
    safe_basename,
    set_custom_properties,
    sha256_file,
)
from lawyer_confirmation_schema import validate_manifest


PRESET = "contract_negotiation_brief"
ASCII_FONT = "Times New Roman"
CHINESE_FONT = "宋体"
BLUE = "000000"
DARK_BLUE = "000000"
INK = "000000"
LIGHT_BLUE = "F2F2F2"
LIGHT_GRAY = "F2F2F2"
MUTED = "666666"
GOLD = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
LABEL_WIDTH_DXA = 2700
VALUE_WIDTH_DXA = TABLE_WIDTH_DXA - LABEL_WIDTH_DXA


def _set_run_font(run, size: float = 11, *, bold: bool = False, color: str = "000000") -> None:
    # Plain law-firm override: Latin characters and numbers use Times New Roman;
    # Word selects SimSun for East Asian glyphs even when both occur in one run.
    run.font.name = ASCII_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), ASCII_FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), ASCII_FONT)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _label_cell(cell, label: str, confirmation_id: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    _set_run_font(run, 9.5, bold=True, color=INK)
    if confirmation_id:
        p.add_run("\n")
        ident = p.add_run(f"Confirmation ID：{confirmation_id}")
        _set_run_font(ident, 8, color=MUTED)


def _plain_sdt(paragraph, *, tag: str, alias: str, initial: str) -> None:
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias_node = OxmlElement("w:alias")
    alias_node.set(qn("w:val"), alias)
    tag_node = OxmlElement("w:tag")
    tag_node.set(qn("w:val"), tag)
    text_node = OxmlElement("w:text")
    sdt_pr.extend([alias_node, tag_node, text_node])
    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    fonts.set(qn("w:hint"), "eastAsia")
    run_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), DARK_BLUE)
    run_pr.append(color)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = initial
    run.extend([run_pr, text])
    sdt_content.append(run)
    sdt.extend([sdt_pr, sdt_content])
    paragraph._p.append(sdt)


def _add_sdt_row(table, *, batch: str, confirmation_id: str, card_id: str, scope: str, field: str,
                 label: str, initial: str, allowed_values: list[str] | None,
                 manifest: list[dict[str, Any]]) -> None:
    cells = table.add_row().cells
    _prevent_row_split(table.rows[-1])
    _label_cell(cells[0], label, confirmation_id)
    tag = f"{batch}/{confirmation_id}/{field}"
    _plain_sdt(cells[1].paragraphs[0], tag=tag, alias=f"{confirmation_id} {label}", initial=initial)
    manifest.append(
        {
            "tag": tag,
            "type": "text",
            "allowed_values": allowed_values,
            "confirmation_id": confirmation_id,
            "card_id": card_id,
            "scope": scope,
            "field": field,
            "editable": True,
        }
    )


def _add_immutable_row(table, *, label: str, value: str) -> None:
    cells = table.add_row().cells
    _prevent_row_split(table.rows[-1])
    _label_cell(cells[0], label)
    cells[1].text = ""
    run = cells[1].paragraphs[0].add_run(value)
    _set_run_font(run, 9.5, color=INK)


def _add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def _add_body(doc, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        _set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        _set_run_font(rest)
    else:
        run = p.add_run(text)
        _set_run_font(run)


def _add_issue_card(doc, issue: dict[str, Any], batch: str, sdt_manifest: list[dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [LABEL_WIDTH_DXA, VALUE_WIDTH_DXA])
    _prevent_row_split(table.rows[0])
    _repeat_table_header(table.rows[0])
    header = table.rows[0].cells
    _label_cell(header[0], f"Issue ID：{issue['issue_id']}")
    _set_cell_shading(header[0], LIGHT_BLUE)
    header[1].text = ""
    p = header[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{issue['file']}｜{issue['clause']}")
    _set_run_font(run, 10, bold=True, color=INK)
    _set_cell_shading(header[1], LIGHT_BLUE)

    immutable_rows = [
        ("风险摘要", issue["risk"]),
        ("法律与处理分析", issue["displayed_analysis"]),
        ("一手依据／来源", "；".join(mapping["clause_locator"] for mapping in issue["source_mappings"])),
        ("Skill 初步方案", f"合同动作：{issue['proposed_drafting_action']}；建议：{issue['displayed_analysis']}"),
        ("同步修改范围", "；".join(issue["sync_scope"])),
    ]
    for label, value in immutable_rows:
        cells = table.add_row().cells
        _prevent_row_split(table.rows[-1])
        _label_cell(cells[0], label)
        cells[1].text = ""
        run = cells[1].paragraphs[0].add_run(value)
        _set_run_font(run, 9.5)

    type_labels = {"fact": "项目事实", "legal": "法律约束", "commercial": "交易选择", "drafting": "文本执行"}
    for item in issue["subitems"]:
        decision_label = (
            f"律师决定｜{type_labels[item['item_type']]}｜"
            f"{item['completion_impact']}｜{'必需' if item['required_for_final'] else '信息'}"
        )
        _add_sdt_row(
            table, batch=batch, confirmation_id=item["confirmation_id"], card_id=issue["issue_id"], scope="subitem",
            field="lawyer_decision", label=decision_label, initial="（请填写）",
            allowed_values=["agree", "revise", "reject", "defer_client", "defer_research", "not_applicable"],
            manifest=sdt_manifest,
        )
        _add_sdt_row(
            table, batch=batch, confirmation_id=item["confirmation_id"], card_id=issue["issue_id"], scope="subitem",
            field="lawyer_comment", label="律师原始意见", initial="（可填写）",
            allowed_values=None, manifest=sdt_manifest,
        )

    issue_id = issue["issue_id"]
    _add_sdt_row(table, batch=batch, confirmation_id=issue_id, card_id=issue_id, scope="issue", field="drafting_action",
                 label="最终合同动作", initial=issue["drafting_action"],
                 allowed_values=["keep_current", "modify", "delete_clause", "no_contract_change", "not_applicable"],
                 manifest=sdt_manifest)
    projection_labels = {
        "client_report_disposition": "客户报告处理",
        "include_in_major_issue_list": "纳入 Major Issue List",
        "include_in_counterparty_comment": "纳入对方批注",
        "include_in_redline": "纳入红线",
    }
    for field in ("client_report_disposition", "include_in_major_issue_list", "include_in_counterparty_comment", "include_in_redline"):
        value = issue["projections"][field]
        initial = str(value).lower() if isinstance(value, bool) else value
        allowed = ["true", "false"] if isinstance(value, bool) else ["include", "client_pending", "legal_pending", "internal_only"]
        _add_sdt_row(table, batch=batch, confirmation_id=issue_id, card_id=issue_id, scope="issue", field=field,
                     label=projection_labels[field], initial=initial, allowed_values=allowed,
                     manifest=sdt_manifest)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_batch_cards(doc, base: dict[str, Any], category: str,
                     sdt_manifest: list[dict[str, Any]]) -> None:
    batches = [item for item in base["batch_decisions"] if item["category"] == category]
    if not batches:
        _add_body(doc, "本批次未设置该类别的批量决定。")
        return
    for batch_item in batches:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _set_table_geometry(table, [LABEL_WIDTH_DXA, VALUE_WIDTH_DXA])
        _prevent_row_split(table.rows[0])
        _repeat_table_header(table.rows[0])
        header = table.rows[0].cells
        _label_cell(header[0], f"Batch ID：{batch_item['batch_id']}")
        _set_cell_shading(header[0], LIGHT_GRAY)
        header[1].text = ""
        run = header[1].paragraphs[0].add_run(
            f"范围：{'、'.join(batch_item['confirmation_ids'])}"
        )
        _set_run_font(run, 9.5, bold=True, color=INK)
        _set_cell_shading(header[1], LIGHT_GRAY)
        _add_sdt_row(
            table,
            batch=base["confirmation_batch_id"],
            confirmation_id=batch_item["batch_id"],
            card_id=batch_item["batch_id"],
            scope="batch",
            field="batch_lawyer_decision",
            label="批量律师决定",
            initial="（请填写）",
            allowed_values=["agree", "revise", "reject", "defer_client", "defer_research", "not_applicable"],
            manifest=sdt_manifest,
        )
        _add_sdt_row(
            table,
            batch=base["confirmation_batch_id"],
            confirmation_id=batch_item["batch_id"],
            card_id=batch_item["batch_id"],
            scope="batch",
            field="batch_exception_confirmation_ids",
            label="例外 Confirmation ID（以逗号或顿号分隔）",
            initial="、".join(batch_item["exception_confirmation_ids"]),
            allowed_values=None,
            manifest=sdt_manifest,
        )
        _add_sdt_row(
            table,
            batch=base["confirmation_batch_id"],
            confirmation_id=batch_item["batch_id"],
            card_id=batch_item["batch_id"],
            scope="batch",
            field="batch_lawyer_comment",
            label="批量律师意见",
            initial="（可填写）",
            allowed_values=None,
            manifest=sdt_manifest,
        )
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    tokens = {
        "Heading 1": (16, BLUE, 14, 8),
        "Heading 2": (13, BLUE, 11, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_document(base: dict[str, Any]) -> tuple[Document, list[dict[str, Any]]]:
    doc = Document()
    _configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("PE/VC 法律审阅｜律师确认单")
    _set_run_font(run, 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{base['matter_id']}｜{base['confirmation_batch_id']}")
    _set_run_font(run, 8, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("审阅关注点确认单")
    _set_run_font(title_run, 21, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("供律师确认｜未经导入验证不得作为最终法律结论")
    _set_run_font(subtitle_run, 10.5, color=GOLD)
    metadata = [
        ("事项编号", base["matter_id"]),
        ("立场", "以基础清单所载审阅立场为准"),
        ("文件版本", base["source_file_collection_sha256"][:16]),
        ("确认批次", base["confirmation_batch_id"]),
        ("审阅轮次", str(base["review_round"])),
    ]
    for label, value in metadata:
        _add_body(doc, f"{label}：{value}", bold_prefix=f"{label}：")

    _add_heading(doc, "1. 使用说明及完成度", 1)
    _add_body(doc, "请在每个内容控件中填写。决定值必须使用表内列示的英文枚举；revise、reject 和 not_applicable 必须说明理由。原生批注只能作为补充。")
    required = [item for issue in base["issues"] for item in issue["subitems"] if item["required_for_final"]]
    facts = [item for issue in base["issues"] for item in issue["subitems"] if item["item_type"] == "fact"]
    drafting = [item for issue in base["issues"] for item in issue["subitems"] if item["item_type"] == "drafting"]
    _add_body(doc, f"完成度（生成时）：决策性必需事项 0/{len(required)}；事实事项 0/{len(facts)}；纯文本清理 0/{len(drafting)}。")

    _add_heading(doc, "2. 项目事实问题", 1)
    for issue in base["issues"]:
        ids = [item["confirmation_id"] for item in issue["subitems"] if item["item_type"] == "fact"]
        if ids:
            _add_body(doc, f"{issue['issue_id']}｜待确认事实：{'、'.join(ids)}。事实决定将在下方同一 Issue 决策卡中填写。")
    if not facts:
        _add_body(doc, "本批次无需要单独向公司追问且会实质改变建议的事实事项。")

    _add_heading(doc, "3. 法律与处理分析", 1)
    _add_body(doc, "每个问题仅设一张决策卡，集中展示事实、法律约束、交易选择和文本执行。")
    sdt_manifest: list[dict[str, Any]] = []
    for issue in base["issues"]:
        doc.add_page_break()
        _add_issue_card(doc, issue, base["confirmation_batch_id"], sdt_manifest)

    _add_heading(doc, "4. 常规实质事项批量确认", 1)
    _add_body(doc, "逐项 Confirmation ID 优先于批量决定；逐项填写后，导入器会将该 ID 自动列入批量例外。")
    _add_batch_cards(doc, base, "regular_substantive", sdt_manifest)
    _add_heading(doc, "5. 纯文本清理批量确认", 1)
    _add_body(doc, "本节仅适用于文本执行事项，不得把事实或法律约束混入纯文本批量确认。")
    _add_batch_cards(doc, base, "text_cleanup", sdt_manifest)

    _add_heading(doc, "6. 律师新增关注点", 1)
    new_table = doc.add_table(rows=1, cols=2)
    new_table.style = "Table Grid"
    _set_table_geometry(new_table, [LABEL_WIDTH_DXA, VALUE_WIDTH_DXA])
    _prevent_row_split(new_table.rows[0])
    _repeat_table_header(new_table.rows[0])
    _label_cell(new_table.rows[0].cells[0], "Issue ID：LAWYER-NEW")
    _set_cell_shading(new_table.rows[0].cells[0], LIGHT_GRAY)
    new_table.rows[0].cells[1].text = "仅本区域可录入新关注点；导入后分配 LAWYER-NEW-### 并进入回读确认。"
    _add_sdt_row(new_table, batch=base["confirmation_batch_id"], confirmation_id="LAWYER-NEW", card_id="LAWYER-NEW", scope="lawyer_new",
                 field="new_issue_text", label="新增关注点", initial="（如有，请填写）",
                 allowed_values=None, manifest=sdt_manifest)

    _add_heading(doc, "7. 未确认、冲突和缺材料摘要", 1)
    unresolved = [item["confirmation_id"] for issue in base["issues"] for item in issue["subitems"] if item["required_for_final"] and item["lawyer_decision"] is None]
    _add_body(doc, f"未确认：{'、'.join(unresolved) if unresolved else '无'}；冲突：生成时未发现；缺材料：以基础清单和源文件集合为准。")

    _add_heading(doc, "8. 律师总体意见、确认声明及日期", 1)
    overall = doc.add_table(rows=0, cols=2)
    overall.style = "Table Grid"
    _set_table_geometry(overall, [LABEL_WIDTH_DXA, VALUE_WIDTH_DXA])
    _add_sdt_row(overall, batch=base["confirmation_batch_id"], confirmation_id="OVERALL", card_id="OVERALL", scope="overall",
                 field="overall_opinion", label="律师总体意见", initial="（可填写）",
                 allowed_values=None, manifest=sdt_manifest)
    _add_sdt_row(
        overall,
        batch=base["confirmation_batch_id"],
        confirmation_id="OVERALL",
        card_id="OVERALL",
        scope="overall",
        field="overall_opinion_effect",
        label="总体意见效力（supplement_only / decision_override）",
        initial="（请填写）",
        allowed_values=["supplement_only", "decision_override"],
        manifest=sdt_manifest,
    )
    _add_immutable_row(
        overall,
        label="确认声明",
        value="总体意见原文仅供审计，其效力以结构化字段为准。本人已审阅上述内容；仅在下栏输入 confirm 后方构成确认。",
    )
    overall_fields = [
        ("confirmation_status", "确认状态（请输入 confirm）", "（请填写）", ["confirm"]),
        ("signoff_name", "确认人", "（请填写）", None),
        ("signoff_date", "确认日期（YYYY-MM-DD）", "（请填写）", None),
    ]
    for field, label, initial, allowed_values in overall_fields:
        _add_sdt_row(overall, batch=base["confirmation_batch_id"], confirmation_id="OVERALL", card_id="OVERALL", scope="overall",
                     field=field, label=label, initial=initial, allowed_values=allowed_values,
                     manifest=sdt_manifest)
    return doc, sorted(sdt_manifest, key=lambda item: item["tag"])


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a Word lawyer-confirmation pack.")
    parser.add_argument("--input", "--base-manifest", dest="base_manifest", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--basename", default="confirmation-form")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    if not safe_basename(args.basename):
        print("unsafe_basename: basename must contain only safe filename characters", file=sys.stderr)
        return 2
    try:
        base = json.loads(args.base_manifest.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        print(f"base_manifest_unreadable: {type(exc).__name__}", file=sys.stderr)
        return 2
    errors = validate_manifest(base)
    if errors:
        print(json.dumps({"code": "base_manifest_invalid", "errors": errors}, ensure_ascii=False), file=sys.stderr)
        return 1
    args.output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = args.output_dir / f"{args.basename}.docx"
    local_base_path = args.output_dir / f"{args.basename}.base-manifest.json"
    pack_path = args.output_dir / f"{args.basename}.pack-manifest.json"
    document, sdt_manifest = build_document(base)
    document.save(docx_path)
    shutil.copyfile(args.base_manifest, local_base_path)
    inspection = inspect_docx(docx_path)
    base_file_sha256 = sha256_file(args.base_manifest)
    sdt_manifest_sha256 = canonical_json_sha256(sdt_manifest)
    generated_content_sha256 = document_part_sha256(docx_path)
    custom_properties = {
        "PEVCGeneratedContentSHA256": generated_content_sha256,
        "PEVCBaseManifestSHA256": base_file_sha256,
        "PEVCImmutableVisibleSHA256": inspection["immutable_visible_content_sha256"],
        "PEVCSDTManifestSHA256": sdt_manifest_sha256,
    }
    set_custom_properties(docx_path, custom_properties)
    pack = {
        "pack_manifest_version": "1.0",
        "design_preset": PRESET,
        "header_template": "memo_masthead",
        "schema_version": base["schema_version"],
        "matter_id": base["matter_id"],
        "review_round": base["review_round"],
        "confirmation_batch_id": base["confirmation_batch_id"],
        "base_manifest_filename": local_base_path.name,
        "base_manifest_file_sha256": base_file_sha256,
        "generated_form_filename": docx_path.name,
        "generated_form_sha256": sha256_file(docx_path),
        "generated_content_sha256": generated_content_sha256,
        "immutable_visible_content_sha256": inspection["immutable_visible_content_sha256"],
        "sdt_manifest_sha256": sdt_manifest_sha256,
        "sdt_manifest": sdt_manifest,
    }
    pack_path.write_text(json.dumps(pack, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    print(json.dumps({"docx": str(docx_path), "base_manifest": str(local_base_path), "pack_manifest": str(pack_path)}, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    sys.exit(main())
