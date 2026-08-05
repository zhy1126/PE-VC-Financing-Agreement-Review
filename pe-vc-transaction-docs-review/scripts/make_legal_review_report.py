#!/usr/bin/env python3
"""Generate a Chinese legal-review DOCX from report-model JSON only."""

from __future__ import annotations

import argparse
from copy import deepcopy
import json
from pathlib import Path
import re
import sys
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_report_model import report_model_content_sha256, validate_report_model_binding


PRESET = "standard_business_brief"
ASCII_FONT = "Times New Roman"
CHINESE_FONT = "宋体"
INK = "000000"
BLUE = "000000"
DARK_BLUE = "000000"
MUTED = "666666"
LIGHT_GRAY = "F2F2F2"
LIGHT_BLUE = "F2F2F2"
CAUTION = "000000"
RISK = "000000"
USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_style_font(style, *, size: float, color: str = "000000", bold: bool | None = None) -> None:
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)
    rfonts.set(qn("w:hint"), "eastAsia")


def _set_run_font(run, *, size: float | None = None, color: str | None = None,
                  bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = ASCII_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)
    rfonts.set(qn("w:hint"), "eastAsia")
    language = rpr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        rpr.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name, size, color in (("Title", 25, INK), ("Subtitle", 12, MUTED)):
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=(name == "Title"))


def _set_cell_margins(cell, *, top: int = 80, bottom: int = 80,
                      start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != USABLE_WIDTH_DXA:
        raise ValueError("table_width_must_equal_9360_dxa")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _paragraph(cell, text: str, *, bold: bool = False, color: str = "000000",
               size: float = 10, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    _set_run_font(p.add_run(text), size=size, color=color, bold=bold)


def _field(paragraph, instruction: str, display: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    _set_run_font(run, size=9, color=MUTED)


def _configure_header_footer(doc: Document, matter_id: str, status: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    _set_run_font(p.add_run(f"法律审查报告  |  {matter_id}"), size=9, color=MUTED, bold=True)
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_table_geometry(table, [6900, 2460], indent=0)
    table.style = "Table Grid"
    # Remove visible grid from the footer while retaining exact geometry.
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    table._tbl.tblPr.append(borders)
    _paragraph(table.cell(0, 0), status, color=MUTED, size=8.5)
    right = table.cell(0, 1)
    right.text = ""
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    _set_run_font(p.add_run("第 "), size=9, color=MUTED)
    _field(p, "PAGE")
    _set_run_font(p.add_run(" 页"), size=9, color=MUTED)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    size, color = {1: (16, BLUE), 2: (13, BLUE), 3: (12, DARK_BLUE)}[level]
    _set_run_font(paragraph.add_run(text), size=size, color=color, bold=True)


def _body(doc: Document, text: str, *, bold_prefix: str | None = None,
          color: str = "000000") -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        _set_run_font(p.add_run(bold_prefix), bold=True, color=color)
        _set_run_font(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        _set_run_font(p.add_run(text), color=color)


def _metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        _paragraph(cells[0], label, bold=True, color=INK, size=9.5)
        _paragraph(cells[1], value, size=9.5)
        _shade(cells[0], LIGHT_GRAY)
    _set_table_geometry(table, [1800, 7560])


def _status_callout(doc: Document, status: str, matter_status: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_geometry(table, [USABLE_WIDTH_DXA])
    _shade(table.cell(0, 0), LIGHT_GRAY)
    _paragraph(table.cell(0, 0), f"文档状态：{status}", bold=True, color=INK, size=11)


def _item_blocks(doc: Document, items: list[dict[str, Any]], *, include_change: bool = False) -> None:
    if not items:
        _body(doc, "无。", color=MUTED)
        return
    for item in items:
        _heading(doc, f"{item['issue_id']}  {item['clause']}", level=2)
        _body(doc, f"风险：{item['risk']}", bold_prefix="风险：")
        if item.get("approved_analysis"):
            _body(doc, f"已批准分析：{item['approved_analysis']}", bold_prefix="已批准分析：")
        if include_change:
            action_label = {
                "modify": "修改",
                "delete_clause": "删除条款并扫描依赖",
                "keep_current": "保留现状",
                "no_contract_change": "不改合同",
                "not_applicable": "不适用",
            }.get(item.get("drafting_action"), str(item.get("drafting_action", "")))
            _body(doc, f"合同动作：{action_label}", bold_prefix="合同动作：")
            sync = "；".join(item.get("sync_scope", [])) or "无"
            _body(doc, f"同步范围：{sync}", bold_prefix="同步范围：")


def _major_issue_table(doc: Document, items: list[dict[str, Any]]) -> None:
    if not items:
        _body(doc, "无。", color=MUTED)
        return
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Issue ID / 条款", "风险摘要", "已批准处理")
    for index, value in enumerate(headers):
        _paragraph(table.rows[0].cells[index], value, bold=True, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(table.rows[0].cells[index], LIGHT_GRAY)
    _repeat_header(table.rows[0])
    for item in items:
        cells = table.add_row().cells
        _paragraph(cells[0], f"{item['issue_id']}\n{item['clause']}", bold=True, color=INK, size=9.2)
        _paragraph(cells[1], item["risk"], size=9.2)
        _paragraph(cells[2], item.get("approved_analysis") or item.get("approved_text") or "无已批准实质内容", size=9.2)
    _set_table_geometry(table, [1700, 2800, 4860])


def _pending_blocks(doc: Document, items: list[dict[str, Any]]) -> None:
    if not items:
        _body(doc, "无。", color=MUTED)
        return
    for item in items:
        _heading(doc, f"{item['issue_id']}  {item['clause']}", level=2)
        _body(doc, f"待决事项：{item.get('pending_reason') or '尚待确认'}", bold_prefix="待决事项：", color=CAUTION)
        _body(doc, "当前处理：不实施实质合同修改；不作为已批准结论。")


def _clean_text(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if re.search(r"(?:^|\s)(?:/" + "Users/" + r"|/home/|[A-Za-z]:\\)", text):
        return Path(text).name
    return text


def _validate_model(model: Any, approved_content: Any) -> dict[str, Any]:
    model = validate_report_model_binding(model, approved_content)
    if not isinstance(model, dict) or model.get("model_type") != "legal_review_report_model":
        raise ValueError("report_model_required")
    required = {
        "model_version", "design_preset", "matter_id", "review_round", "confirmation_batch_id",
        "matter_status", "visible_status", "position", "scope", "version_basis", "limitations",
        "unreviewed_materials", "projections", "anti_leak", "required_section_ids", "report_model_sha256",
    }
    if not required.issubset(model):
        raise ValueError("report_model_incomplete")
    if model["design_preset"] != PRESET:
        raise ValueError("standard_business_brief_required")
    if model["matter_status"] not in {"blocked", "passed_with_limitations", "passed"}:
        raise ValueError("invalid_matter_status")
    if model["report_model_sha256"] != report_model_content_sha256(model):
        raise ValueError("report_model_hash_mismatch")
    client = model.get("projections", {}).get("client_report")
    major = model.get("projections", {}).get("major_issue_list")
    if not isinstance(client, dict) or not isinstance(major, dict):
        raise ValueError("client_and_major_projections_required")
    client_anti_leak = model.get("anti_leak", {}).get("client_report", {})
    excluded = set(client_anti_leak.get("conclusion_excluded_ids", []))
    allowed_pending = set(client_anti_leak.get("allowed_pending_ids", []))
    included = set(client.get("included_ids", []))
    pending = set(client.get("client_pending_ids", [])) | set(client.get("legal_pending_ids", []))
    if excluded & included:
        raise ValueError("client_report_anti_leak_overlap")
    if pending != allowed_pending or not pending.issubset(excluded):
        raise ValueError("client_pending_anti_leak_mismatch")
    return deepcopy(model)


def make_legal_review_report(report_model: Any, approved_content: Any, output: Path, *,
                             final: bool = True) -> Path:
    model = _validate_model(report_model, approved_content)
    status = model["matter_status"]
    if final and status == "blocked":
        raise ValueError("blocked_cannot_be_final")
    visible_status = {
        "passed": "最终版",
        "passed_with_limitations": "最终版（附保留事项）",
        "blocked": "草稿——律师确认未完成",
    }[status]
    client = model["projections"]["client_report"]
    major = model["projections"]["major_issue_list"]
    included = sorted(client.get("items", []), key=lambda value: value["issue_id"])

    doc = Document()
    _configure_document(doc)
    _configure_header_footer(doc, _clean_text(model["matter_id"]), visible_status)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(title.add_run("法律审查报告"), size=22, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_run_font(subtitle.add_run("PE/VC 交易文件律师确认版"), size=12, color=MUTED)
    _status_callout(doc, visible_status, status)
    _metadata_table(doc, [
        ("事项编号", _clean_text(model["matter_id"])),
        ("审阅轮次", str(model["review_round"])),
        ("确认批次", _clean_text(model["confirmation_batch_id"])),
        ("审阅立场", _clean_text(model["position"])),
    ])

    _heading(doc, "范围、版本、立场与限制")
    scope = "；".join(Path(_clean_text(item)).name for item in model["scope"]) or "未列明"
    _body(doc, f"审阅范围：{scope}", bold_prefix="审阅范围：")
    _body(doc, f"版本基础：{_clean_text(model['version_basis'])}", bold_prefix="版本基础：")
    _body(doc, f"审阅立场：{_clean_text(model['position'])}", bold_prefix="审阅立场：")
    for limitation in model["limitations"]:
        _body(doc, f"限制：{_clean_text(limitation)}", bold_prefix="限制：", color=CAUTION if status != "passed" else MUTED)

    _heading(doc, "执行摘要")
    _body(doc, f"文档状态：{visible_status}。客户报告纳入 {len(client.get('included_ids', []))} 项；Major Issue List 纳入 {len(major.get('included_ids', []))} 项。")
    if status == "passed_with_limitations":
        _body(doc, "本报告仅对局部待决事项作显著保留，不表示相关文件已达到交付签署条件。", color=CAUTION)
    elif status == "blocked":
        _body(doc, "律师确认或回读尚未完成，本文件仅供内部继续审阅。", color=RISK)

    _heading(doc, "Major Issue List")
    _major_issue_table(doc, sorted(major.get("items", []), key=lambda value: value["issue_id"]))

    _heading(doc, "已确认事实基础")
    if included:
        for item in included:
            _body(doc, f"{item['issue_id']}｜{Path(_clean_text(item['file'])).name}｜{item['clause']}：以律师已确认的文件事实及范围为分析基础。")
    else:
        _body(doc, "无可纳入的已确认事实基础。", color=MUTED)

    _heading(doc, "已批准法律与处理分析")
    _item_blocks(doc, [item for item in included if item.get("approved_analysis")])

    _heading(doc, "修改建议及同步修改")
    _item_blocks(doc, [item for item in included if item.get("drafting_action") in {"modify", "delete_clause"}], include_change=True)

    _heading(doc, "待客户确认事项")
    _pending_blocks(doc, sorted(client.get("client_pending_items", []), key=lambda value: value["issue_id"]))

    _heading(doc, "待进一步法律核验事项")
    _pending_blocks(doc, sorted(client.get("legal_pending_items", []), key=lambda value: value["issue_id"]))

    _heading(doc, "未审材料")
    materials = [_clean_text(item) for item in model["unreviewed_materials"] if _clean_text(item)]
    if materials:
        for material in materials:
            _body(doc, material)
    else:
        _body(doc, "无另行列明的未审材料。", color=MUTED)

    doc.core_properties.title = "法律审查报告"
    doc.core_properties.subject = visible_status
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def _load_model(path: Path) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise ValueError("report_model_required") from exc


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a legal-review DOCX from report-model JSON only.")
    parser.add_argument("--report-model", required=True, type=Path)
    parser.add_argument("--approved-content", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--draft", action="store_true", help="Allow a clearly labeled blocked draft.")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    try:
        make_legal_review_report(
            _load_model(args.report_model), _load_model(args.approved_content),
            args.output, final=not args.draft,
        )
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 1
    print(json.dumps({"status": "generated", "output": args.output.name, "design_preset": PRESET}, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
